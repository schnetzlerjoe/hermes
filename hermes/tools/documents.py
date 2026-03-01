"""Document generation tools for creating Word reports and exporting to PDF.

Built on python-docx for Word documents.  PDF export uses LibreOffice headless
CLI for best fidelity -- it renders the .docx exactly as Word would, including
tables, images, and page breaks.
"""

from __future__ import annotations

import logging
import subprocess
import uuid
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from llama_index.core.tools import FunctionTool

from hermes.config import get_config

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# In-memory document store
# ---------------------------------------------------------------------------

# Maps document IDs to python-docx Document objects.
_documents: dict[str, Document] = {}


def _get_document(doc_id: str) -> Document:
    """Retrieve an open document or raise a clear error."""
    if doc_id not in _documents:
        raise ValueError(
            f"Document '{doc_id}' not found.  Available: {list(_documents.keys())}"
        )
    return _documents[doc_id]


# ---------------------------------------------------------------------------
# Tool functions
# ---------------------------------------------------------------------------


def doc_create(name: str, title: str | None = None) -> str:
    """Create a new Word document.

    If a *title* is provided it is added as a centred ``Title`` paragraph
    at the top of the document.

    Args:
        name: Human-readable name used as the document identifier.
        title: Optional document title to insert on the first page.

    Returns:
        Document ID string for use in subsequent operations.
    """
    doc = Document()
    doc_id = f"{name}_{uuid.uuid4().hex[:8]}"

    if title:
        p = doc.add_paragraph(title, style="Title")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _documents[doc_id] = doc
    logger.info("Created document %s", doc_id)
    return f"Created document '{doc_id}'."


def doc_add_heading(doc_id: str, text: str, level: int = 1) -> str:
    """Add a heading to the document.

    Args:
        doc_id: Document ID.
        text: Heading text.
        level: Heading level (1-4).  Level 1 is the largest.

    Returns:
        Confirmation string.
    """
    doc = _get_document(doc_id)
    level = max(1, min(level, 4))
    doc.add_heading(text, level=level)
    return f"Added heading level {level}: '{text}'."


def doc_add_paragraph(
    doc_id: str,
    text: str,
    style: str | None = None,
    bold: bool = False,
    italic: bool = False,
) -> str:
    """Add a paragraph with optional styling.

    Args:
        doc_id: Document ID.
        text: Paragraph text.
        style: Optional Word style name (e.g. ``"List Bullet"``,
            ``"Intense Quote"``).
        bold: Make the entire paragraph bold.
        italic: Make the entire paragraph italic.

    Returns:
        Confirmation string.
    """
    doc = _get_document(doc_id)
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return f"Added paragraph ({len(text)} chars)."


def doc_add_table(
    doc_id: str,
    headers: list[str],
    rows: list[list[str]],
    title: str | None = None,
    style: str = "Light Grid Accent 1",
) -> str:
    """Insert a formatted table, optionally preceded by a bold title paragraph.

    The *title* is written to the document immediately before the table in the
    same call, guaranteeing the label always appears above the data regardless
    of parallel tool-call execution order.

    The table is set to stretch to 100 % of the page text-width so it never
    overflows the margins.

    Args:
        doc_id: Document ID.
        headers: Column header labels.
        rows: 2D list of string values (one inner list per row).
        title: Optional bold paragraph inserted directly above the table.
        style: Word table style name.

    Returns:
        Confirmation string with table dimensions.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = _get_document(doc_id)

    # Optional title paragraph — written before the table so order is guaranteed.
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True

    n_cols = len(headers)
    n_rows = len(rows) + 1  # +1 for header

    table = doc.add_table(rows=n_rows, cols=n_cols)

    # Apply style -- fall back gracefully if the style name is invalid.
    try:
        table.style = style
    except KeyError:
        pass  # Default style is acceptable.

    # Stretch table to 100 % of the page text-width so columns never overflow.
    tblPr = table._tbl.tblPr
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")   # 5000 / 100 = 100 %
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)

    # Write headers with bold formatting.
    for col_idx, header in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Write data rows.
    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            if col_idx < n_cols:
                table.rows[row_idx + 1].cells[col_idx].text = str(value)

    return f"Added {len(rows)}x{n_cols} table."


def doc_add_image(
    doc_id: str,
    image_path: str,
    width_inches: float = 6.0,
) -> str:
    """Insert an image into the document.

    Typically used to embed chart PNGs generated by the charts tools.

    Args:
        doc_id: Document ID.
        image_path: Absolute or relative path to the image file.
        width_inches: Display width in inches.  Height is scaled
            proportionally.

    Returns:
        Confirmation string.
    """
    doc = _get_document(doc_id)
    image_path_obj = Path(image_path)
    if not image_path_obj.exists():
        raise FileNotFoundError(f"Image not found: {image_path}")

    doc.add_picture(str(image_path_obj), width=Inches(width_inches))
    return f"Added image '{image_path_obj.name}' ({width_inches}\" wide)."


def doc_add_page_break(doc_id: str) -> str:
    """Insert a page break.

    Args:
        doc_id: Document ID.

    Returns:
        Confirmation string.
    """
    doc = _get_document(doc_id)
    doc.add_page_break()
    return "Page break inserted."


def doc_save(doc_id: str, filename: str | None = None) -> str:
    """Save the document as a ``.docx`` file.

    Saves to the configured output directory.  The document remains open
    in memory for further modifications.

    Args:
        doc_id: Document ID.
        filename: Optional filename (without path).  Defaults to
            ``"{doc_id}.docx"``.

    Returns:
        Absolute path to the saved file.
    """
    doc = _get_document(doc_id)
    cfg = get_config()

    output_dir = Path(cfg.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    if filename is None:
        filename = f"{doc_id}.docx"
    if not filename.endswith(".docx"):
        filename += ".docx"

    filepath = output_dir / filename
    doc.save(str(filepath))
    logger.info("Saved document to %s", filepath.resolve())
    return str(filepath.resolve())


def doc_export_pdf(
    docx_path: str,
    output_dir: str | None = None,
) -> str:
    """Convert a ``.docx`` file to PDF using LibreOffice headless.

    LibreOffice must be installed on the system (``libreoffice`` or
    ``soffice`` in PATH).  This produces the highest-fidelity conversion
    available without Microsoft Word.

    Args:
        docx_path: Path to the ``.docx`` file to convert.
        output_dir: Optional output directory for the PDF.  Defaults to
            the same directory as the input file.

    Returns:
        Absolute path to the generated PDF file.

    Raises:
        FileNotFoundError: If the docx file does not exist.
        RuntimeError: If LibreOffice is not installed or conversion fails.
    """
    docx_file = Path(docx_path)
    if not docx_file.exists():
        raise FileNotFoundError(f"File not found: {docx_path}")

    if output_dir is None:
        target_dir = docx_file.parent
    else:
        target_dir = Path(output_dir)
        target_dir.mkdir(parents=True, exist_ok=True)

    try:
        result = subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(target_dir),
                str(docx_file),
            ],
            capture_output=True,
            text=True,
            timeout=120,
        )
    except FileNotFoundError:
        raise RuntimeError(
            "LibreOffice not found.  Install it with: "
            "apt-get install libreoffice-writer (Debian/Ubuntu) or "
            "brew install --cask libreoffice (macOS)."
        ) from None

    if result.returncode != 0:
        raise RuntimeError(
            f"LibreOffice conversion failed (exit {result.returncode}): "
            f"{result.stderr.strip()}"
        )

    pdf_path = target_dir / docx_file.with_suffix(".pdf").name
    logger.info("Exporting PDF: %s", pdf_path)
    if not pdf_path.exists():
        raise RuntimeError(
            f"Conversion appeared to succeed but PDF not found at {pdf_path}. "
            f"LibreOffice stdout: {result.stdout.strip()}"
        )

    return str(pdf_path.resolve())


def doc_read(doc_id: str) -> str:
    """Read back the full text content of an in-memory document for review.

    Returns a structured text representation of all headings, paragraphs,
    and tables so the agent can verify content before saving.  Every block
    is prefixed with its 1-based body index ``[N]`` so you can pass it to
    ``doc_edit_paragraph``.  Tables are also numbered ``[TABLE T]`` for use
    with ``doc_edit_table_cell``.

    Args:
        doc_id: Document ID returned by doc_create.

    Returns:
        Full document content as a structured string.
    """
    from docx.oxml.ns import qn

    _WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    doc = _get_document(doc_id)
    lines: list[str] = []
    table_count = 0

    for block_idx, block in enumerate(doc.element.body, start=1):
        tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag

        if tag == "p":
            style_el = block.find(qn("w:pPr"))
            style_name = ""
            if style_el is not None:
                style_id_el = style_el.find(qn("w:pStyle"))
                if style_id_el is not None:
                    style_name = style_id_el.get(qn("w:val"), "")
            text = "".join(r.text for r in block.iter() if r.tag == qn("w:t") and r.text)
            if not text.strip():
                continue
            if style_name.startswith("Heading"):
                level = style_name[-1] if style_name[-1].isdigit() else "1"
                lines.append(f"[{block_idx}] [HEADING {level}] {text}")
            else:
                lines.append(f"[{block_idx}] [PARA] {text}")

        elif tag == "tbl":
            table_count += 1
            rows_text: list[list[str]] = []
            for row in block.iter(f"{{{_WNS}}}tr"):
                cells = []
                for cell in row.iter(f"{{{_WNS}}}tc"):
                    cell_text = "".join(
                        t.text for t in cell.iter(f"{{{_WNS}}}t") if t.text
                    )
                    cells.append(cell_text.strip())
                if cells:
                    rows_text.append(cells)
            if rows_text:
                lines.append(
                    f"[{block_idx}] [TABLE {table_count}: "
                    f"{len(rows_text)}x{len(rows_text[0])}]"
                )
                for row_i, row in enumerate(rows_text, start=1):
                    lines.append(f"  [r{row_i}] " + " | ".join(row))

    return "\n".join(lines) if lines else "(empty document)"


def doc_edit_paragraph(doc_id: str, block_index: int, new_text: str) -> str:
    """Replace the text of an existing body paragraph or heading.

    Use ``doc_read`` first to identify the correct ``block_index`` (the ``[N]``
    prefix shown for each block).  Only paragraph/heading blocks can be edited
    with this tool — for table cells use ``doc_edit_table_cell``.

    Args:
        doc_id: Document ID.
        block_index: 1-based body-block index as shown by doc_read.
        new_text: Replacement text (preserves the paragraph's style).

    Returns:
        Confirmation string.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = _get_document(doc_id)
    blocks = list(doc.element.body)
    if block_index < 1 or block_index > len(blocks):
        raise ValueError(
            f"block_index {block_index} out of range (1..{len(blocks)}). "
            "Use doc_read to list valid indices."
        )
    block = blocks[block_index - 1]
    tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag
    if tag != "p":
        raise ValueError(
            f"Block {block_index} is a '{tag}', not a paragraph. "
            "Use doc_edit_table_cell for table content."
        )
    # Remove all existing runs, keeping paragraph properties (style) intact.
    for run_el in block.findall(qn("w:r")):
        block.remove(run_el)
    # Insert a single new run with the replacement text.
    r_el = OxmlElement("w:r")
    t_el = OxmlElement("w:t")
    t_el.text = new_text
    if new_text and (new_text[0] == " " or new_text[-1] == " "):
        t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r_el.append(t_el)
    block.append(r_el)
    return f"Block {block_index} updated ({len(new_text)} chars)."


def doc_edit_table_cell(
    doc_id: str,
    table_index: int,
    row_index: int,
    col_index: int,
    new_text: str,
) -> str:
    """Replace the text of a single table cell.

    Use ``doc_read`` to identify the table number (``[TABLE T: ...]``) and row
    number (``[rN]``).  Column index is 1-based left-to-right.

    Args:
        doc_id: Document ID.
        table_index: 1-based table number as shown by doc_read (TABLE 1, TABLE 2, …).
        row_index: 1-based row number (1 = header row, 2 = first data row, …).
        col_index: 1-based column number.
        new_text: Replacement cell text.

    Returns:
        Confirmation string.
    """
    doc = _get_document(doc_id)
    tables = doc.tables
    if table_index < 1 or table_index > len(tables):
        raise ValueError(
            f"table_index {table_index} out of range (1..{len(tables)}). "
            "Use doc_read to see available tables."
        )
    table = tables[table_index - 1]
    if row_index < 1 or row_index > len(table.rows):
        raise ValueError(
            f"row_index {row_index} out of range (1..{len(table.rows)})."
        )
    row = table.rows[row_index - 1]
    if col_index < 1 or col_index > len(row.cells):
        raise ValueError(
            f"col_index {col_index} out of range (1..{len(row.cells)})."
        )
    row.cells[col_index - 1].text = new_text
    return f"Table {table_index} cell [r{row_index}, c{col_index}] updated to {new_text!r}."


# ---------------------------------------------------------------------------
# Tool registration
# ---------------------------------------------------------------------------


def create_tools() -> list[FunctionTool]:
    """Create LlamaIndex FunctionTool instances for all document tools."""
    return [
        FunctionTool.from_defaults(
            fn=doc_create,
            name="doc_create",
            description=(
                "Create a new Word document with an optional title. "
                "Returns a document ID for subsequent operations."
            ),
        ),
        FunctionTool.from_defaults(
            fn=doc_add_heading,
            name="doc_add_heading",
            description="Add a heading (levels 1-4) to a Word document.",
        ),
        FunctionTool.from_defaults(
            fn=doc_add_paragraph,
            name="doc_add_paragraph",
            description=(
                "Add a paragraph to a Word document with optional bold, "
                "italic, and style settings."
            ),
        ),
        FunctionTool.from_defaults(
            fn=doc_add_table,
            name="doc_add_table",
            description=(
                "Insert a formatted table with headers and rows into a "
                "Word document."
            ),
        ),
        FunctionTool.from_defaults(
            fn=doc_add_image,
            name="doc_add_image",
            description=(
                "Insert an image (PNG chart, etc.) into a Word document "
                "at a specified width."
            ),
        ),
        FunctionTool.from_defaults(
            fn=doc_add_page_break,
            name="doc_add_page_break",
            description="Insert a page break into a Word document.",
        ),
        FunctionTool.from_defaults(
            fn=doc_read,
            name="doc_read",
            description=(
                "Read back the full content of an in-memory document as structured text. "
                "Each block is prefixed with its [N] index for use with doc_edit_paragraph. "
                "Tables are numbered [TABLE T] for use with doc_edit_table_cell."
            ),
        ),
        FunctionTool.from_defaults(
            fn=doc_edit_paragraph,
            name="doc_edit_paragraph",
            description=(
                "Replace the text of an existing paragraph or heading by its block index "
                "(shown as [N] in doc_read output). Use to fix mistakes without adding "
                "new paragraphs."
            ),
        ),
        FunctionTool.from_defaults(
            fn=doc_edit_table_cell,
            name="doc_edit_table_cell",
            description=(
                "Replace the text of a single table cell. Specify table number (TABLE T "
                "from doc_read), 1-based row and column. Use to correct individual cell "
                "values without re-writing the whole table."
            ),
        ),
        FunctionTool.from_defaults(
            fn=doc_save,
            name="doc_save",
            description=(
                "Save a Word document as .docx. Returns the absolute file path."
            ),
        ),
        FunctionTool.from_defaults(
            fn=doc_export_pdf,
            name="doc_export_pdf",
            description=(
                "Convert a .docx file to PDF using LibreOffice headless. "
                "Returns the absolute path to the generated PDF."
            ),
        ),
    ]
