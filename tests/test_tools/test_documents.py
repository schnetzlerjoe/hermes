"""Tests for document generation tools.

All tests create real .docx files in temporary directories and verify
their contents using python-docx.  No mocking of document functionality.
"""

from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt

# ---------------------------------------------------------------------------
# Helpers -- these exercise the same operations the document tools will
# provide.  We test the underlying python-docx behaviour so that the
# tests validate real file I/O.
# ---------------------------------------------------------------------------


def create_document(title: str | None = None) -> Document:
    """Create a new Word document, optionally with a title heading."""
    doc = Document()
    if title:
        doc.add_heading(title, level=0)
    return doc


class TestCreateDocument:
    """Test basic document creation."""

    def test_creates_empty_document(self, tmp_output_dir: Path) -> None:
        """An empty document should save as a valid .docx file."""
        path = tmp_output_dir / "test_empty.docx"
        doc = Document()
        doc.save(str(path))

        assert path.exists()
        assert path.stat().st_size > 0

    def test_document_with_title(self, tmp_output_dir: Path) -> None:
        """A document created with a title should contain that heading."""
        path = tmp_output_dir / "test_title.docx"
        doc = create_document("Equity Research Report")
        doc.save(str(path))

        doc2 = Document(str(path))
        # The title is a Heading 0 paragraph.
        assert len(doc2.paragraphs) >= 1
        assert doc2.paragraphs[0].text == "Equity Research Report"


class TestAddHeading:
    """Test adding headings at various levels."""

    def test_heading_level_1(self, tmp_output_dir: Path) -> None:
        """A level-1 heading should persist with correct text."""
        path = tmp_output_dir / "test_h1.docx"
        doc = Document()
        doc.add_heading("Business Overview", level=1)
        doc.save(str(path))

        doc2 = Document(str(path))
        heading = doc2.paragraphs[0]
        assert heading.text == "Business Overview"
        assert heading.style.name == "Heading 1"

    def test_heading_level_2(self, tmp_output_dir: Path) -> None:
        """A level-2 heading should use the Heading 2 style."""
        path = tmp_output_dir / "test_h2.docx"
        doc = Document()
        doc.add_heading("Revenue Analysis", level=2)
        doc.save(str(path))

        doc2 = Document(str(path))
        heading = doc2.paragraphs[0]
        assert heading.text == "Revenue Analysis"
        assert heading.style.name == "Heading 2"

    def test_multiple_headings_in_order(self, tmp_output_dir: Path) -> None:
        """Multiple headings should appear in insertion order."""
        path = tmp_output_dir / "test_multi_heading.docx"
        doc = Document()
        sections = [
            "Executive Summary",
            "Financial Analysis",
            "Valuation",
            "Risks",
            "Recommendation",
        ]
        for section in sections:
            doc.add_heading(section, level=1)
        doc.save(str(path))

        doc2 = Document(str(path))
        headings = [p.text for p in doc2.paragraphs if p.style.name.startswith("Heading")]
        assert headings == sections


class TestAddParagraph:
    """Test adding body text paragraphs."""

    def test_simple_paragraph(self, tmp_output_dir: Path) -> None:
        """A paragraph should preserve its full text content."""
        path = tmp_output_dir / "test_paragraph.docx"
        text = (
            "Apple Inc. reported revenue of $394.3 billion for fiscal year 2024, "
            "representing a 2% year-over-year increase driven by strong Services growth."
        )
        doc = Document()
        doc.add_paragraph(text)
        doc.save(str(path))

        doc2 = Document(str(path))
        assert doc2.paragraphs[0].text == text

    def test_bold_and_italic_runs(self, tmp_output_dir: Path) -> None:
        """Bold and italic formatting within a paragraph should persist."""
        path = tmp_output_dir / "test_runs.docx"
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("We rate ")
        run_bold = para.add_run("AAPL")
        run_bold.bold = True
        para.add_run(" as ")
        run_italic = para.add_run("Overweight")
        run_italic.italic = True
        para.add_run(".")
        doc.save(str(path))

        doc2 = Document(str(path))
        para2 = doc2.paragraphs[0]
        assert para2.text == "We rate AAPL as Overweight."

        # Verify individual run formatting.
        runs = para2.runs
        assert any(r.bold for r in runs)
        assert any(r.italic for r in runs)

    def test_paragraph_with_font_size(self, tmp_output_dir: Path) -> None:
        """Custom font size should be preserved."""
        path = tmp_output_dir / "test_font_size.docx"
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Disclaimer: This is not investment advice.")
        run.font.size = Pt(8)
        doc.save(str(path))

        doc2 = Document(str(path))
        run2 = doc2.paragraphs[0].runs[0]
        assert run2.font.size == Pt(8)


class TestAddTable:
    """Test adding tables to documents."""

    def test_simple_table(self, tmp_output_dir: Path) -> None:
        """A table should preserve its row and column count."""
        path = tmp_output_dir / "test_table.docx"
        doc = Document()

        headers = ["Metric", "2022", "2023", "2024"]
        rows = [
            ["Revenue ($B)", "394.3", "383.3", "390.0"],
            ["Net Income ($B)", "99.8", "97.0", "93.7"],
            ["EPS ($)", "6.15", "6.13", "6.08"],
        ]

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"

        # Write header row.
        for idx, header in enumerate(headers):
            table.rows[0].cells[idx].text = header

        # Write data rows.
        for row_data in rows:
            row_cells = table.add_row().cells
            for idx, value in enumerate(row_data):
                row_cells[idx].text = value

        doc.save(str(path))

        doc2 = Document(str(path))
        table2 = doc2.tables[0]
        assert len(table2.rows) == 4  # 1 header + 3 data rows
        assert len(table2.columns) == 4

    def test_table_cell_values(self, tmp_output_dir: Path) -> None:
        """Individual cell values should be readable after save/load."""
        path = tmp_output_dir / "test_table_values.docx"
        doc = Document()

        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Ticker"
        table.cell(0, 1).text = "Price"
        table.cell(1, 0).text = "AAPL"
        table.cell(1, 1).text = "$195.50"
        doc.save(str(path))

        doc2 = Document(str(path))
        t = doc2.tables[0]
        assert t.cell(0, 0).text == "Ticker"
        assert t.cell(0, 1).text == "Price"
        assert t.cell(1, 0).text == "AAPL"
        assert t.cell(1, 1).text == "$195.50"

    def test_table_with_merged_cells(self, tmp_output_dir: Path) -> None:
        """Merged cells should survive save/load."""
        path = tmp_output_dir / "test_merged.docx"
        doc = Document()

        table = doc.add_table(rows=3, cols=3)
        # Merge the first row across all columns for a section header.
        merged_cell = table.cell(0, 0).merge(table.cell(0, 2))
        merged_cell.text = "Financial Highlights"

        table.cell(1, 0).text = "Revenue"
        table.cell(1, 1).text = "394.3"
        table.cell(1, 2).text = "383.3"

        doc.save(str(path))

        doc2 = Document(str(path))
        t = doc2.tables[0]
        assert t.cell(0, 0).text == "Financial Highlights"


class TestSaveToDisk:
    """Test document save and file integrity."""

    def test_save_creates_file(self, tmp_output_dir: Path) -> None:
        """Saving should produce a non-empty file on disk."""
        path = tmp_output_dir / "test_save.docx"
        doc = Document()
        doc.add_paragraph("Test content")
        doc.save(str(path))

        assert path.exists()
        assert path.stat().st_size > 0

    def test_saved_docx_is_valid_zip(self, tmp_output_dir: Path) -> None:
        """A .docx file is a ZIP archive; verify valid structure."""
        path = tmp_output_dir / "test_docx_zip.docx"
        doc = Document()
        doc.add_paragraph("ZIP check")
        doc.save(str(path))

        assert zipfile.is_zipfile(str(path))

    def test_saved_docx_contains_content_types(self, tmp_output_dir: Path) -> None:
        """A valid .docx must contain [Content_Types].xml."""
        path = tmp_output_dir / "test_content_types.docx"
        doc = Document()
        doc.add_paragraph("Content types check")
        doc.save(str(path))

        with zipfile.ZipFile(str(path)) as zf:
            names = zf.namelist()
            assert "[Content_Types].xml" in names

    def test_full_report_structure(self, tmp_output_dir: Path) -> None:
        """A document with headings, paragraphs, and tables should all persist."""
        path = tmp_output_dir / "test_full_report.docx"
        doc = Document()

        doc.add_heading("Equity Research: AAPL", level=0)
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(
            "Apple remains well-positioned with strong ecosystem lock-in "
            "and growing Services revenue."
        )
        doc.add_heading("Financial Summary", level=1)

        table = doc.add_table(rows=2, cols=3)
        table.style = "Table Grid"
        for idx, header in enumerate(["Metric", "FY2023", "FY2024"]):
            table.rows[0].cells[idx].text = header
        for idx, val in enumerate(["Revenue", "$383.3B", "$394.3B"]):
            table.rows[1].cells[idx].text = val

        doc.add_heading("Recommendation", level=1)
        doc.add_paragraph("Buy with a 12-month price target of $220.")

        doc.save(str(path))

        # Reload and verify structure.
        doc2 = Document(str(path))
        assert len(doc2.paragraphs) >= 4  # headings + paragraphs
        assert len(doc2.tables) == 1
        assert doc2.tables[0].cell(1, 0).text == "Revenue"
